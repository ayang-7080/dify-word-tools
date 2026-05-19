import json
from collections.abc import Generator
from io import BytesIO
from typing import Any

from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage
from docx import Document

from tools.tools import no_proxy


class AyToolsTool(Tool):
    def _invoke(self, tool_parameters: dict[str, Any]) -> Generator[ToolInvokeMessage]:
        file = tool_parameters.get('file')
        comments = tool_parameters.get('comments')

        try:
            comments = json.loads(comments)
        except Exception:
            raise ValueError('comments must be a JSON object')
        filename = file.filename
        #print(filename)
        with no_proxy():
            file_data = file.blob

        #print("下载word成功")
        doc = Document(BytesIO(file_data))
        doc = add_comments_to_document(doc, comments)
        #print("处理成功")

        byte_io = BytesIO()
        doc.save(byte_io)
        byte_io.seek(0)
        docx_bytes = byte_io.getvalue()  # 使用 getvalue() 替代 read()
        byte_io.close()  # 显式关闭

        # 返回 blob
        yield self.create_blob_message(
            blob=docx_bytes,
            meta={
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "filename": f"{filename}"  # 可选：指定文件名
            }
        )



def add_comments_to_document(doc , comment_map):
    """
    comment_map: { "待查找文本": "批注内容" }
    """
    for para in doc.paragraphs:
        for target, comment_text in comment_map.items():
            _process_paragraph(doc, para, target, comment_text)

    return doc


def _process_paragraph(doc, para, target, comment_text):
    """
    在一个段落中循环查找并添加批注
    """
    # 采用倒序查找，防止拆分 Run 后索引偏移影响后续匹配
    # 这里简单起见使用 full_text.find，如果需要正则可用 re.finditer
    while True:
        full_text = "".join(run.text for run in para.runs)
        start_idx = full_text.find(target)
        if start_idx == -1:
            break

        # 核心：将包含 target 的部分拆分成一个独立的 Run
        target_run = _isolate_target_run(para, start_idx, len(target))

        # 使用官方原生 API 添加批注
        doc.add_comment(
            runs=target_run,
            text=comment_text,
            author="AI-Audit",
            initials="AA"
        )

        # 为了避免死循环，这里可以临时改名或记录已处理位置
        # 这里演示简单处理：只处理第一个匹配项，或在逻辑中标记
        # 实际生产建议记录位置偏移，此处为演示精简逻辑
        break


def _isolate_target_run(para, start_idx, length):
    """
    精准拆分 Run，将 target 文本隔离到一个单独的新 Run 中并返回
    """
    current_pos = 0
    target_run = None

    for run in list(para.runs):
        run_len = len(run.text)
        if current_pos <= start_idx < current_pos + run_len:
            # 找到起始位置所在的 Run
            rel_start = start_idx - current_pos
            rel_end = rel_start + length

            full_run_text = run.text

            # 1. 处理起始 Run 之前的文本
            run.text = full_run_text[:rel_start]

            # 2. 创建匹配文本的 Run，并复制原 Run 格式
            matched_run = para.add_run(full_run_text[rel_start:rel_end])
            _copy_run_format(run, matched_run)

            # 3. 处理匹配文本之后的文本
            after_text = full_run_text[rel_end:]
            if after_text:
                after_run = para.add_run(after_text)
                _copy_run_format(run, after_run)

            # 4. 将新生成的 Run 移动到正确位置（add_run 默认在末尾）
            _move_run_after(run, matched_run)
            if after_text:
                _move_run_after(matched_run, after_run)

            return matched_run

        current_pos += run_len
    return None


def _copy_run_format(src_run, dst_run):
    """复制 Run 的基础格式"""
    dst_run.bold = src_run.bold
    dst_run.italic = src_run.italic
    dst_run.underline = src_run.underline
    dst_run.font.name = src_run.font.name
    dst_run.font.size = src_run.font.size
    if src_run.font.color.rgb:
        dst_run.font.color.rgb = src_run.font.color.rgb


def _move_run_after(ref_run, new_run):
    """底层 XML 操作：将 new_run 移动到 ref_run 后面"""
    ref_run._element.addnext(new_run._element)
