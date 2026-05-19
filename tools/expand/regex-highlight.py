from collections.abc import Generator
from io import BytesIO
from typing import Any

from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage

import re

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor

from tools.tools import no_proxy


class AyToolsTool(Tool):
    def _invoke(self, tool_parameters: dict[str, Any]) -> Generator[ToolInvokeMessage]:
        file = tool_parameters.get('file')
        content_type = tool_parameters.get('content_type')
        in_regex = tool_parameters.get('in_regex',"")

        is_italic = tool_parameters.get('is_italic',False)
        is_bold = tool_parameters.get('is_bold',False)

        with no_proxy():
            file_data = file.blob

        filename = file.filename

        doc = Document(BytesIO(file_data))

        if content_type == "pmid":
            pattern = r"[（\(][^）\)]*?(?:PMID[:：]\s*\d+|10\.\d{4,9}/[^\s,;；、\)\]\}\）】]+)(?:\s*[;；]\s*(?:PMID[:：]\s*\d+|10\.\d{4,9}/[^\s,;；、\)\]\}\）】]+))*\s*[）\)]"
        elif content_type == "image_number":
            pattern = r"[（\(]\s*(?:(?:图|附图|表|Figure|Figures?|Table)\s*(?:[Ss]?\d+)?(?:[A-Z](?:[–-][A-Z])?)?(?:\s*[,，]\s*(?:(?:图|附图|表|Figure|Figures?|Table)\s*(?:[Ss]?\d+)?(?:[A-Z](?:[–-][A-Z])?)?))*)\s*[）\)]"
        elif content_type == "pmid_and_image":
            pattern = r"[（\(][^）\)]*?(?:PMID[:：]\s*\d+|10\.\d{4,9}/[^\s,;；、\)\]\}\）】]+)(?:\s*[;；]\s*(?:PMID[:：]\s*\d+|10\.\d{4,9}/[^\s,;；、\)\]\}\）】]+))*\s*[）\)]"
            pattern2 = r"[（\(]\s*(?:(?:图|附图|表|Figure|Figures?|Table)\s*(?:[Ss]?\d+)?(?:[A-Z](?:[–-][A-Z])?)?(?:\s*[,，]\s*(?:(?:图|附图|表|Figure|Figures?|Table)\s*(?:[Ss]?\d+)?(?:[A-Z](?:[–-][A-Z])?)?))*)\s*[）\)]"
        elif content_type == "customize_regex":
            pattern = in_regex
        else:
            raise Exception("请选择高亮内容类型")
        if pattern == "" or pattern is None:
            raise Exception("缺少正则表达式")

        for paragraph in doc.paragraphs:
            #print(paragraph.text)
            highlight_across_runs(
                paragraph,
                pattern,
                color=RGBColor(0, 0, 255),  # 蓝色
                italic=is_italic,
                bold=is_bold,
            )
            if content_type == "pmid_and_image":
                highlight_across_runs(
                    paragraph,
                    pattern2,
                    color=RGBColor(0, 0, 255),  # 蓝色
                    italic=is_italic,
                    bold=is_bold,
                )
        byte_io = BytesIO()
        doc.save(byte_io)
        byte_io.seek(0)
        docx_bytes = byte_io.getvalue()  # 使用 getvalue() 替代 read()
        byte_io.close()  # 显式关闭

        # 返回 blob
        yield self.create_blob_message(
            blob=docx_bytes,
            meta={
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "filename": f"{filename}"  # 可选：指定文件名
            }
        )



def insert_run_after(paragraph, ref_run, new_text):
    """
    在 ref_run 后面插入一个新的 Run（保持位置正确）
    """
    new_run = paragraph.add_run(new_text)
    # 获取 XML 元素
    ref_run_element = ref_run._r
    new_run_element = new_run._r

    # 在 ref_run_element 后面插入 new_run_element
    parent = ref_run_element.getparent()
    if parent is not None:
        index = parent.index(ref_run_element)
        parent.insert(index + 1, new_run_element)
    return new_run


def copy_run_format(source_run, target_run):
    """完整复制格式，特别强化中英文字体设置"""
    font_src = source_run.font
    font_tgt = target_run.font

    # 基础格式
    font_tgt.bold = font_src.bold
    font_tgt.italic = font_src.italic
    font_tgt.underline = font_src.underline
    font_tgt.size = font_src.size

    # 🎯 关键：显式设置英文字体和中文字体
    try:
        if font_src.name:
            font_tgt.name = font_src.name  # 设置西文字体
            # 设置中文字体（东亚字体）
            if hasattr(target_run._element, 'rPr') and target_run._element.rPr is not None:
                rPr = target_run._element.rPr
                if not hasattr(rPr, 'rFonts'):
                    from docx.oxml import OxmlElement
                    rFonts = OxmlElement('w:rFonts')
                    rPr.append(rFonts)
                else:
                    rFonts = rPr.rFonts
                rFonts.set(qn('w:eastAsia'), font_src.name)  # 强制设置中文字体
                rFonts.set(qn('w:ascii'), font_src.name)     # 可选：设置英文字体
                rFonts.set(qn('w:hAnsi'), font_src.name)     # 可选：设置 ANSI 字体
    except Exception as e:
        pass
    # 颜色（在设置字体后再设置颜色，避免触发重置）
    if font_src.color and font_src.color.rgb:
        font_tgt.color.rgb = font_src.color.rgb

def highlight_across_runs(paragraph, pattern, color=None, italic=False, bold=False):
    """
    支持跨多个 Run 匹配并高亮（解决括号被拆到不同 Run 的问题）
    """
    if isinstance(pattern, str):
        pattern = re.compile(pattern)

    runs = paragraph.runs
    if not runs:
        return

    # Step 1: 合并所有 Run 的文本，记录每个字符属于哪个 Run 和位置
    full_text = ""
    char_to_run_index = []  # 记录每个字符属于哪个 run
    run_start_indices = []  # 每个 run 在 full_text 中的起始位置

    for i, run in enumerate(runs):
        run_text = run.text
        start_index = len(full_text)
        full_text += run_text
        run_start_indices.append(start_index)
        char_to_run_index.extend([i] * len(run_text))

    # Step 2: 在完整文本中查找所有匹配项
    matches = list(pattern.finditer(full_text))
    if not matches:
        return

    # Step 3: 从后往前处理每个匹配（避免索引偏移）
    for match in reversed(matches):
        start_pos = match.start()
        end_pos = match.end()
        matched_text = match.group()

        # 找到匹配范围涉及哪些 Run
        start_run_idx = char_to_run_index[start_pos]
        end_run_idx = char_to_run_index[end_pos - 1]  # end_pos 是开区间

        # 如果匹配跨越多个 Run，我们需要合并或逐段处理
        if start_run_idx == end_run_idx:
            # 简单情况：匹配在一个 Run 内
            run = runs[start_run_idx]
            local_start = start_pos - run_start_indices[start_run_idx]
            local_end = end_pos - run_start_indices[start_run_idx]

            before_text = run.text[:local_start]
            target_text = run.text[local_start:local_end]
            after_text = run.text[local_end:]

            run.text = before_text
            new_run = insert_run_after(paragraph, run, target_text)
            copy_run_format(run, new_run)
            if color:
                new_run.font.color.rgb = color
            if italic:
                new_run.font.italic = True
            if bold:
                new_run.font.bold = True
            if after_text:
                after_run = insert_run_after(paragraph, new_run, after_text)
                copy_run_format(run, after_run)

        else:
            # 复杂情况：匹配跨越多个 Run
            # 我们把涉及的所有 Run 合并成一个临时文本，处理后再拆分
            combined_text = ""
            involved_runs = []
            start_in_first_run = start_pos - run_start_indices[start_run_idx]

            for i in range(start_run_idx, end_run_idx + 1):
                involved_runs.append(runs[i])
                combined_text += runs[i].text

            # 计算在 combined_text 中的位置
            local_start = start_in_first_run
            local_end = end_pos - run_start_indices[start_run_idx]

            before_text = combined_text[:local_start]
            target_text = combined_text[local_start:local_end]
            after_text = combined_text[local_end:]

            # 清空第一个 Run，写入 before_text
            first_run = involved_runs[0]
            first_run.text = before_text

            # 插入高亮部分
            new_run = insert_run_after(paragraph, first_run, target_text)
            copy_run_format(first_run, new_run)
            if color:
                new_run.font.color.rgb = color
            if italic:
                new_run.font.italic = True
            if bold:
                new_run.font.bold = True

            # 处理剩余文本：如果 after_text 非空，插入到 new_run 后
            if after_text:
                after_run = insert_run_after(paragraph, new_run, after_text)
                copy_run_format(first_run, after_run)

            # 删除中间和最后一个 Run（因为内容已合并处理）
            for i in range(start_run_idx + 1, end_run_idx + 1):
                runs[i].text = ""  # 清空，避免残留

    return True
