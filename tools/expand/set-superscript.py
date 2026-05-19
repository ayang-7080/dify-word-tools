import re
from io import BytesIO
from typing import Any
from collections.abc import Generator
from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn

from tools.tools import no_proxy


class AyToolsTool(Tool):
    def _invoke(self, tool_parameters: dict[str, Any]) -> Generator[ToolInvokeMessage]:
        file = tool_parameters.get('file')
        content_type = tool_parameters.get('content_type')
        in_regex = tool_parameters.get('in_regex', "")

        # 用户自定义颜色，默认为蓝色
        color_hex = tool_parameters.get('color', "#0000FF")
        parsed_color = self._parse_hex_color(color_hex)
        with no_proxy():
            file_data = file.blob
        filename = file.filename
        doc = Document(BytesIO(file_data))

        # --- 设置匹配模式 ---
        if content_type == "citation":
            # 匹配 [1], [1-2], [1][2], [1,2,3] 等模式
            pattern = r'\[\d+(?:[–-]\d+)?\](?:\[\d+(?:[–-]\d+)?\])*|\[\d+(?:\s*,\s*\d+)*\]'
        elif content_type == "customize_regex":
            pattern = in_regex
        else:
            # 沿用你之前的逻辑
            pattern = in_regex

        if not pattern:
            raise Exception("缺少正则表达式或匹配类型")

        # 遍历文档处理
        for paragraph in doc.paragraphs:
            set_superscript_across_runs(
                paragraph,
                pattern,
                color=parsed_color
            )

        # 保存并返回
        byte_io = BytesIO()
        doc.save(byte_io)
        byte_io.seek(0)
        yield self.create_blob_message(
            blob=byte_io.getvalue(),
            meta={
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "filename": f"{filename}"
            }
        )

    def _parse_hex_color(self, hex_str):
        hex_str = hex_str.lstrip('#')
        return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))


# --- 核心处理函数 ---

def set_superscript_across_runs(paragraph, pattern, color=None):
    """
    匹配文本并设置为上标，同时可选设置颜色
    """
    if isinstance(pattern, str):
        pattern = re.compile(pattern)

    runs = paragraph.runs
    if not runs: return

    # Step 1: 建立文本映射（逻辑同你提供的代码）
    full_text = ""
    char_to_run_index = []
    run_start_indices = []

    for i, run in enumerate(runs):
        run_text = run.text
        start_index = len(full_text)
        full_text += run_text
        run_start_indices.append(start_index)
        char_to_run_index.extend([i] * len(run_text))

    matches = list(pattern.finditer(full_text))
    if not matches: return
    #print(matches)
    # Step 2: 逆序处理匹配项
    for match in reversed(matches):
        start_pos = match.start()
        end_pos = match.end()
        matched_text = match.group()

        start_run_idx = char_to_run_index[start_pos]
        end_run_idx = char_to_run_index[end_pos - 1]

        if start_run_idx == end_run_idx:
            # --- 情况 A: 匹配在单个 Run 内 ---
            run = runs[start_run_idx]
            local_start = start_pos - run_start_indices[start_run_idx]
            local_end = end_pos - run_start_indices[start_run_idx]

            before_text = run.text[:local_start]
            target_text = run.text[local_start:local_end]
            after_text = run.text[local_end:]

            run.text = before_text
            new_run = insert_run_after(paragraph, run, target_text)
            copy_run_format(run, new_run)

            # ✨ 设置关键属性
            new_run.font.superscript = True
            if color:
                new_run.font.color.rgb = color

            if after_text:
                after_run = insert_run_after(paragraph, new_run, after_text)
                copy_run_format(run, after_run)

        else:
            # --- 情况 B: 跨 Run 匹配 ---
            combined_text = ""
            involved_runs = []
            start_in_first_run = start_pos - run_start_indices[start_run_idx]

            for i in range(start_run_idx, end_run_idx + 1):
                involved_runs.append(runs[i])
                combined_text += runs[i].text

            local_start = start_in_first_run
            local_end = end_pos - run_start_indices[start_run_idx]

            before_text = combined_text[:local_start]
            target_text = combined_text[local_start:local_end]
            after_text = combined_text[local_end:]

            first_run = involved_runs[0]
            first_run.text = before_text

            new_run = insert_run_after(paragraph, first_run, target_text)
            copy_run_format(first_run, new_run)

            # ✨ 设置关键属性
            new_run.font.superscript = True
            if color:
                new_run.font.color.rgb = color

            if after_text:
                after_run = insert_run_after(paragraph, new_run, after_text)
                copy_run_format(first_run, after_run)

            for i in range(start_run_idx + 1, end_run_idx + 1):
                runs[i].text = ""

    return True


# 复用你提供的辅助函数
def insert_run_after(paragraph, ref_run, new_text):
    new_run = paragraph.add_run(new_text)
    ref_run_element = ref_run._r
    new_run_element = new_run._r
    parent = ref_run_element.getparent()
    if parent is not None:
        index = parent.index(ref_run_element)
        parent.insert(index + 1, new_run_element)
    return new_run


def copy_run_format(source_run, target_run):
    font_src = source_run.font
    font_tgt = target_run.font
    font_tgt.bold = font_src.bold
    font_tgt.italic = font_src.italic
    font_tgt.underline = font_src.underline
    font_tgt.size = font_src.size
    try:
        if font_src.name:
            font_tgt.name = font_src.name
            rPr = target_run._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:eastAsia'), font_src.name)
            rFonts.set(qn('w:ascii'), font_src.name)
            rFonts.set(qn('w:hAnsi'), font_src.name)
    except:
        pass
    if font_src.color and font_src.color.rgb:
        font_tgt.color.rgb = font_src.color.rgb