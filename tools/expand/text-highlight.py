from collections.abc import Generator
from io import BytesIO
from typing import Any

from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage

import re

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor

class AyToolsTool(Tool):
    def _invoke(self, tool_parameters: dict[str, Any]) -> Generator[ToolInvokeMessage]:
        file = tool_parameters.get('file')
        texts = tool_parameters.get('texts')
        font_color = tool_parameters.get('font_color',None)
        italic = tool_parameters.get('italic',False)
        bold = tool_parameters.get('bold', False)
        if font_color == "":
            font_color = None
        if italic == "":
            italic = False
        if bold == "":
            bold = False

        file_data = file.blob

        filename = file.filename

        doc = Document(BytesIO(file_data))

        keywords = [kw.strip() for kw in re.split(r'[\n;；,，]', texts) if kw.strip()]

        for paragraph in doc.paragraphs:
            highlight_text_list(
                paragraph,
                keywords,
                color = font_color,
                italic = italic,
                bold = bold
            )

        byte_io = BytesIO()
        doc.save(byte_io)
        byte_io.seek(0)
        docx_bytes = byte_io.getvalue()  # 使用 getvalue() 替代 read()
        byte_io.close()  # 显式关闭

        # 返回 blob
        yield self.create_blob_message(
            blob=docx_bytes,
            meta={
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "filename": f"{filename}"  # 可选：指定文件名
            }
        )



def insert_run_after(paragraph, ref_run, new_text):
    """
    在 ref_run 后面插入一个新的 Run（保持位置正确）
    """
    new_run = paragraph.add_run(new_text)
    # 获取 XML 元素
    ref_run_element = ref_run._r
    new_run_element = new_run._r

    # 在 ref_run_element 后面插入 new_run_element
    parent = ref_run_element.getparent()
    if parent is not None:
        index = parent.index(ref_run_element)
        parent.insert(index + 1, new_run_element)
    return new_run


def copy_run_format(source_run, target_run):
    """完整复制格式，特别强化中英文字体设置"""
    font_src = source_run.font
    font_tgt = target_run.font

    # 基础格式
    font_tgt.bold = font_src.bold
    font_tgt.italic = font_src.italic
    font_tgt.underline = font_src.underline
    font_tgt.size = font_src.size

    # 🎯 关键：显式设置英文字体和中文字体
    try:
        if font_src.name:
            font_tgt.name = font_src.name  # 设置西文字体
            # 设置中文字体（东亚字体）
            if hasattr(target_run._element, 'rPr') and target_run._element.rPr is not None:
                rPr = target_run._element.rPr
                if not hasattr(rPr, 'rFonts'):
                    from docx.oxml import OxmlElement
                    rFonts = OxmlElement('w:rFonts')
                    rPr.append(rFonts)
                else:
                    rFonts = rPr.rFonts
                rFonts.set(qn('w:eastAsia'), font_src.name)  # 强制设置中文字体
                rFonts.set(qn('w:ascii'), font_src.name)     # 可选：设置英文字体
                rFonts.set(qn('w:hAnsi'), font_src.name)     # 可选：设置 ANSI 字体
    except Exception as e:
        pass
    # 颜色（在设置字体后再设置颜色，避免触发重置）
    if font_src.color and font_src.color.rgb:
        font_tgt.color.rgb = font_src.color.rgb



def parse_color(color_input):
    """
    支持十六进制颜色字符串（如 "#FF0000"）或 RGBColor 对象
    """
    if isinstance(color_input, str) and color_input.startswith('#'):
        # 去掉 #，转为 RGB
        hex_color = color_input.lstrip('#')
        if len(hex_color) == 6:
            r = int(hex_color[0:2], 16)
            g = int(hex_color[2:4], 16)
            b = int(hex_color[4:6], 16)
            return RGBColor(r, g, b)
        else:
            raise ValueError(f"Invalid hex color: {color_input}")
    return color_input  # 假设已经是 RGBColor 对象

def highlight_text_list(paragraph, text_list, color=None, italic=False, bold=False):
    """
    在段落中查找并高亮指定的普通文本（支持字符串列表）
    不依赖正则，精确匹配
    兼容现有函数调用方式，不修改 highlight_across_runs
    """
    if not text_list:
        return

    # 如果是单个字符串，转为列表
    if isinstance(text_list, str):
        text_list = [text_list]

    runs = paragraph.runs
    if not runs:
        return

    # Step 1: 合并所有 Run 文本，记录映射
    full_text = ""
    char_to_run_index = []
    run_start_indices = []

    for i, run in enumerate(runs):
        run_text = run.text
        start_index = len(full_text)
        full_text += run_text
        run_start_indices.append(start_index)
        char_to_run_index.extend([i] * len(run_text))

    # Step 2: 对每个目标文本进行匹配
    for target_text in text_list:
        if not target_text or len(target_text) == 0:
            continue

        start_pos = 0
        while True:
            # 查找下一个匹配位置
            pos = full_text.find(target_text, start_pos)
            if pos == -1:
                break

            end_pos = pos + len(target_text)

            # 找到涉及的 Run
            start_run_idx = char_to_run_index[pos]
            end_run_idx = char_to_run_index[end_pos - 1]

            if start_run_idx == end_run_idx:
                # 单 Run 内
                run = runs[start_run_idx]
                local_start = pos - run_start_indices[start_run_idx]
                local_end = end_pos - run_start_indices[start_run_idx]

                before_text = run.text[:local_start]
                matched_text = run.text[local_start:local_end]
                after_text = run.text[local_end:]

                run.text = before_text
                new_run = insert_run_after(paragraph, run, matched_text)
                copy_run_format(run, new_run)
                parsed_color = parse_color(color)
                if parsed_color:
                    new_run.font.color.rgb = parsed_color
                if italic:
                    new_run.font.italic = True
                if bold:
                    new_run.font.bold = True
                if after_text:
                    after_run = insert_run_after(paragraph, new_run, after_text)
                    copy_run_format(run, after_run)

            else:
                # 跨 Run
                combined_text = ""
                involved_runs = []
                start_in_first_run = pos - run_start_indices[start_run_idx]

                for i in range(start_run_idx, end_run_idx + 1):
                    involved_runs.append(runs[i])
                    combined_text += runs[i].text

                local_start = start_in_first_run
                local_end = end_pos - run_start_indices[start_run_idx]

                before_text = combined_text[:local_start]
                matched_text = combined_text[local_start:local_end]
                after_text = combined_text[local_end:]

                first_run = involved_runs[0]
                first_run.text = before_text

                new_run = insert_run_after(paragraph, first_run, matched_text)
                copy_run_format(first_run, new_run)
                parsed_color = parse_color(color)
                if parsed_color:
                    new_run.font.color.rgb = parsed_color
                if italic:
                    new_run.font.italic = True
                if bold:
                    new_run.font.bold = True

                if after_text:
                    after_run = insert_run_after(paragraph, new_run, after_text)
                    copy_run_format(first_run, after_run)

                for i in range(start_run_idx + 1, end_run_idx + 1):
                    runs[i].text = ""

            # 继续查找下一个匹配（避免重叠）
            start_pos = end_pos

    return True

def highlight_across_runs(paragraph, pattern, color=None, italic=False, bold=False):
    """
    支持跨多个 Run 匹配并高亮（解决括号被拆到不同 Run 的问题）
    """
    if isinstance(pattern, str):
        pattern = re.compile(pattern)

    runs = paragraph.runs
    if not runs:
        return

    # Step 1: 合并所有 Run 的文本，记录每个字符属于哪个 Run 和位置
    full_text = ""
    char_to_run_index = []  # 记录每个字符属于哪个 run
    run_start_indices = []  # 每个 run 在 full_text 中的起始位置

    for i, run in enumerate(runs):
        run_text = run.text
        start_index = len(full_text)
        full_text += run_text
        run_start_indices.append(start_index)
        char_to_run_index.extend([i] * len(run_text))

    # Step 2: 在完整文本中查找所有匹配项
    matches = list(pattern.finditer(full_text))
    if not matches:
        return

    # Step 3: 从后往前处理每个匹配（避免索引偏移）
    for match in reversed(matches):
        start_pos = match.start()
        end_pos = match.end()
        matched_text = match.group()

        # 找到匹配范围涉及哪些 Run
        start_run_idx = char_to_run_index[start_pos]
        end_run_idx = char_to_run_index[end_pos - 1]  # end_pos 是开区间

        # 如果匹配跨越多个 Run，我们需要合并或逐段处理
        if start_run_idx == end_run_idx:
            # 简单情况：匹配在一个 Run 内
            run = runs[start_run_idx]
            local_start = start_pos - run_start_indices[start_run_idx]
            local_end = end_pos - run_start_indices[start_run_idx]

            before_text = run.text[:local_start]
            target_text = run.text[local_start:local_end]
            after_text = run.text[local_end:]

            run.text = before_text
            new_run = insert_run_after(paragraph, run, target_text)
            copy_run_format(run, new_run)
            if color:
                new_run.font.color.rgb = color
            if italic:
                new_run.font.italic = True
            if after_text:
                after_run = insert_run_after(paragraph, new_run, after_text)
                copy_run_format(run, after_run)

        else:
            # 复杂情况：匹配跨越多个 Run
            # 我们把涉及的所有 Run 合并成一个临时文本，处理后再拆分
            combined_text = ""
            involved_runs = []
            start_in_first_run = start_pos - run_start_indices[start_run_idx]

            for i in range(start_run_idx, end_run_idx + 1):
                involved_runs.append(runs[i])
                combined_text += runs[i].text

            # 计算在 combined_text 中的位置
            local_start = start_in_first_run
            local_end = end_pos - run_start_indices[start_run_idx]

            before_text = combined_text[:local_start]
            target_text = combined_text[local_start:local_end]
            after_text = combined_text[local_end:]

            # 清空第一个 Run，写入 before_text
            first_run = involved_runs[0]
            first_run.text = before_text

            # 插入高亮部分
            new_run = insert_run_after(paragraph, first_run, target_text)
            copy_run_format(first_run, new_run)
            if color:
                new_run.font.color.rgb = color
            if italic:
                new_run.font.italic = True

            # 处理剩余文本：如果 after_text 非空，插入到 new_run 后
            if after_text:
                after_run = insert_run_after(paragraph, new_run, after_text)
                copy_run_format(first_run, after_run)

            # 删除中间和最后一个 Run（因为内容已合并处理）
            for i in range(start_run_idx + 1, end_run_idx + 1):
                runs[i].text = ""  # 清空，避免残留

    return True
