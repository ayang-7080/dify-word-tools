import os
import time
import traceback
from contextlib import contextmanager
from typing import Optional

import httpx
from datetime import datetime
import random
import zipfile
import xml.etree.ElementTree as ET
import base64
import re
from io import BytesIO

from PIL import Image
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn, nsdecls

from collections import Counter
def generate_timestamp_random():
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    random_suffix = random.randint(1000, 9999)
    return f"{timestamp}{random_suffix}"

def check_proxy():
    resp = httpx.get("http://httpbin.org/ip")
    origin =  resp.json()['origin']
    print("当前IP:", resp.json()['origin'])
    return origin

def extract_docx_to_json(docx_path,table_markdown = False):
    result = []

    # 命名空间
    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    }

    style_font_size_map = {}  # 样式 -> 字号
    style_align_map = {}  # 样式 -> 对齐方式
    style_bold_map = {}  # --- 新增：样式是否加粗映射 ---

    with zipfile.ZipFile(docx_path) as docx_zip:

        # 提取 styles.xml，构建样式ID -> 字号 的映射
        if "word/styles.xml" in docx_zip.namelist():
            styles_xml = ET.fromstring(docx_zip.read("word/styles.xml"))
            for style in styles_xml.findall("w:style", ns):
                style_id = style.attrib.get(f"{{{ns['w']}}}styleId")
                rPr = style.find("w:rPr", ns)
                if rPr is not None:
                    sz = rPr.find("w:sz", ns)
                    if sz is not None:
                        size_val = int(sz.attrib.get(f"{{{ns['w']}}}val")) / 2
                        style_font_size_map[style_id] = size_val
                    # --- 新增：样式级加粗检测 ---
                    b_tag = rPr.find("w:b", ns)
                    # 如果有 <w:b /> 且 w:val 不为 0/false/off，则视为加粗
                    if b_tag is not None:
                        val = b_tag.attrib.get(f"{{{ns['w']}}}val", "true")
                        style_bold_map[style_id] = val.lower() not in ["0", "false", "off"]
                pPr = style.find("w:pPr", ns)
                if pPr is not None:
                    jc = pPr.find("w:jc", ns)
                    if jc is not None:
                        style_align_map[style_id] = jc.attrib.get(f"{{{ns['w']}}}val", "left")

        # 读取文档主内容
        document_xml = docx_zip.read("word/document.xml")
        root = ET.fromstring(document_xml)

        # 读取图片路径映射（关系文件）
        rels_map = {}
        rels_path = "word/_rels/document.xml.rels"
        if rels_path in docx_zip.namelist():
            rels_xml = ET.fromstring(docx_zip.read(rels_path))
            for rel in rels_xml.findall("Relationship",
                                        {"": "http://schemas.openxmlformats.org/package/2006/relationships"}):
                r_id = rel.attrib.get("Id")
                target = rel.attrib.get("Target")
                if r_id and target and target.startswith("media/"):
                    rels_map[r_id] = target

        # 遍历 body 子节点，按顺序识别段落、表格、图片
        body = root.find("w:body", ns)
        for child in body:
            tag = child.tag.split("}")[-1]

            if tag == "p":  # 段落，可能包含文本或图片
                text = "".join([t.text for t in child.findall(".//w:t", ns) if t.text])
                if text.strip():
                    pPr = child.find("w:pPr", ns)
                    alignment = "left"
                    first_line_indent = 0
                    font_size = None
                    font_name = "宋体"

                    para_style_id = None

                    if pPr is not None:

                        # 样式 ID
                        style_tag = pPr.find("w:pStyle", ns)
                        if style_tag is not None:
                            para_style_id = style_tag.attrib.get(f"{{{ns['w']}}}val")

                        # 对齐方式
                        jc = pPr.find("w:jc", ns)
                        if jc is not None:
                            alignment = jc.attrib.get(f"{{{ns['w']}}}val", "left")
                        elif para_style_id and para_style_id in style_align_map:
                            alignment = style_align_map[para_style_id]

                        # 首行缩进
                        ind = pPr.find("w:ind", ns)
                        if ind is not None and f"{{{ns['w']}}}firstLine" in ind.attrib:
                            first_line_indent = int(ind.attrib.get(f"{{{ns['w']}}}firstLine")) / 20  # 转 pt

                    font_sizes = []
                    bold_flags = []  # 记录各 run 的加粗情况

                    for r in child.findall("w:r", ns):
                        rPr = r.find("w:rPr", ns)
                        is_r_bold = False
                        if rPr is not None:
                            # 字号
                            sz = rPr.find("w:sz", ns)
                            if sz is not None:
                                sz_val = int(sz.attrib.get(f"{{{ns['w']}}}val")) / 2
                                font_sizes.append(sz_val)

                            # --- 新增：运行级加粗检测 ---
                            b_tag = rPr.find("w:b", ns)
                            if b_tag is not None:
                                val = b_tag.attrib.get(f"{{{ns['w']}}}val", "true")
                                is_r_bold = val.lower() not in ["0", "false", "off"]
                            elif para_style_id in style_bold_map:
                                is_r_bold = style_bold_map[para_style_id]

                            # 字体
                            fonts = rPr.find("w:rFonts", ns)
                            if fonts is not None:
                                font_name = fonts.attrib.get(f"{{{ns['w']}}}ascii", font_name)
                        else:
                            # 如果没有 rPr，则看段落样式是否加粗
                            if para_style_id in style_bold_map:
                                is_r_bold = style_bold_map[para_style_id]

                        bold_flags.append(is_r_bold)

                    # print(text.strip())
                    # print(font_sizes)
                    # 根据最多出现的字号选取
                    if font_sizes:
                        counter = Counter(font_sizes)
                        most_common = counter.most_common()
                        if most_common:  # 再次判断是否非空，确保稳健
                            max_freq = most_common[0][1]
                            # 取所有“频率等于最大频率”的字号
                            most_frequent_sizes = [size for size, freq in most_common if freq == max_freq]
                            # 从这些中选最大字号
                            font_size = max(most_frequent_sizes)
                        else:
                            font_size = 12
                    elif para_style_id and para_style_id in style_font_size_map:
                        font_size = style_font_size_map[para_style_id]
                    else:
                        font_size = 12
                    # print(font_size)

                    # --- 改进后的判定逻辑：必须超过 2/3 的内容加粗才视为整段加粗 ---
                    is_bold = False
                    if bold_flags:
                        try:
                            true_count = bold_flags.count(True)
                            total_count = len(bold_flags)
                            if total_count == 2: # 需求 1：如果是 2 个块，只要有 1 个加粗（50%）就视为整行加粗
                                is_bold = (true_count >= 1)
                            elif true_count / total_count > 2 / 3: # 判断 True 的比例是否超过 2/3 (约 66.7%)
                                is_bold = True
                            else:
                                is_bold = False
                        except Exception as e:
                            is_bold = False

                    # 拼接到结果中
                    result.append({
                        "type": "text",
                        "content": text.strip(),
                        "style": {
                            "alignment": alignment,
                            "first_line_indent_pt": first_line_indent,
                            "font_name": font_name,
                            "font_size_pt": font_size,
                            "bold":is_bold
                        }
                    })
                    #print("*" * 30)

                # 查找 embedded 图片
                for blip in child.findall(".//a:blip", ns):
                    r_embed = blip.attrib.get("{%s}embed" % ns["r"])
                    if r_embed and r_embed in rels_map:
                        image_path = "word/" + rels_map[r_embed]
                        if image_path in docx_zip.namelist():
                            image_data = docx_zip.read(image_path)
                            b64 = base64.b64encode(image_data).decode()

                            # 获取图片大小
                            width_pt = height_pt = width_cm = height_cm = None
                            for drawing in child.findall(".//w:drawing", ns):
                                extent = drawing.find(".//wp:extent", {
                                    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                                })
                                if extent is not None:
                                    cx = int(extent.attrib.get("cx", 0))
                                    cy = int(extent.attrib.get("cy", 0))
                                    width_pt = round(cx / 12700, 2)  # 转 pt
                                    height_pt = round(cy / 12700, 2)
                                    width_cm = round(cx / 360000, 2)
                                    height_cm = round(cy / 360000, 2)
                                    break  # 找到一个即可

                            result.append({
                                "type": "image",
                                "filename": os.path.basename(image_path),
                                "content": b64,
                                "size": {
                                    "width_pt": width_pt,
                                    "height_pt": height_pt,
                                    "width_cm": width_cm,
                                    "height_cm": height_cm
                                }
                            })

            elif tag == "tbl2":  # 表格
                line_strip = "\n"
                if table_markdown:
                    line_strip = "<br>"
                rows = []
                for row in child.findall("w:tr", ns):
                    cells = []
                    for cell in row.findall("w:tc", ns):
                        texts = [t.text for t in cell.findall(".//w:t", ns) if t.text]
                        cells.append(line_strip.join(texts).strip())
                    rows.append(cells)

                if table_markdown:
                    try:
                        rows = convert_table_to_markdown(rows)
                    except Exception as e:
                        pass
                result.append({"type": "table", "content": rows})


            elif tag == "tbl":  # 表格
                line_strip = "\n"
                if table_markdown:
                    line_strip = "<br>"
                rows = []
                for row in child.findall("w:tr", ns):
                    cells = []
                    for cell in row.findall("w:tc", ns):
                        cell_lines = []
                        # 逐段落解析（w:p）
                        for p in cell.findall("w:p", ns):
                            line_parts = []
                            for node in p.iter():
                                tag_name = node.tag.split("}")[-1]
                                # 文本节点
                                if tag_name == "t" and node.text:
                                    line_parts.append(node.text)
                                # 软换行 <w:br/> 或 <w:cr/>
                                elif tag_name in ("br", "cr"):
                                    line_parts.append(line_strip)
                            line = "".join(line_parts).strip()
                            if line:
                                cell_lines.append(line)
                        # 用换行拼接每个段落
                        cell_text = line_strip.join(cell_lines).strip()
                        cells.append(cell_text)
                    rows.append(cells)

                if table_markdown:

                    try:

                        rows = convert_table_to_markdown(rows)

                    except Exception:

                        pass

                result.append({"type": "table", "content": rows})

    return result

#
# def extract_docx_to_json(docx_path,table_markdown = False):
#     result = []
#
#     # 命名空间
#     ns = {
#         "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
#         "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
#         "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
#     }
#
#     style_font_size_map = {}  # 样式 -> 字号
#     style_align_map = {}  # 样式 -> 对齐方式
#
#     with zipfile.ZipFile(docx_path) as docx_zip:
#
#         # 提取 styles.xml，构建样式ID -> 字号 的映射
#         if "word/styles.xml" in docx_zip.namelist():
#             styles_xml = ET.fromstring(docx_zip.read("word/styles.xml"))
#             for style in styles_xml.findall("w:style", ns):
#                 style_id = style.attrib.get(f"{{{ns['w']}}}styleId")
#                 rPr = style.find("w:rPr", ns)
#                 if rPr is not None:
#                     sz = rPr.find("w:sz", ns)
#                     if sz is not None:
#                         size_val = int(sz.attrib.get(f"{{{ns['w']}}}val")) / 2
#                         style_font_size_map[style_id] = size_val
#                 pPr = style.find("w:pPr", ns)
#                 if pPr is not None:
#                     jc = pPr.find("w:jc", ns)
#                     if jc is not None:
#                         style_align_map[style_id] = jc.attrib.get(f"{{{ns['w']}}}val", "left")
#
#         # 读取文档主内容
#         document_xml = docx_zip.read("word/document.xml")
#         root = ET.fromstring(document_xml)
#
#         # 读取图片路径映射（关系文件）
#         rels_map = {}
#         rels_path = "word/_rels/document.xml.rels"
#         if rels_path in docx_zip.namelist():
#             rels_xml = ET.fromstring(docx_zip.read(rels_path))
#             for rel in rels_xml.findall("Relationship",
#                                         {"": "http://schemas.openxmlformats.org/package/2006/relationships"}):
#                 r_id = rel.attrib.get("Id")
#                 target = rel.attrib.get("Target")
#                 if r_id and target and target.startswith("media/"):
#                     rels_map[r_id] = target
#
#         # 遍历 body 子节点，按顺序识别段落、表格、图片
#         body = root.find("w:body", ns)
#         for child in body:
#             tag = child.tag.split("}")[-1]
#
#             if tag == "p":  # 段落，可能包含文本或图片
#                 text = "".join([t.text for t in child.findall(".//w:t", ns) if t.text])
#                 if text.strip():
#                     pPr = child.find("w:pPr", ns)
#                     alignment = "left"
#                     first_line_indent = 0
#                     font_size = None
#                     font_name = "宋体"
#
#                     para_style_id = None
#
#                     if pPr is not None:
#
#                         # 样式 ID
#                         style_tag = pPr.find("w:pStyle", ns)
#                         if style_tag is not None:
#                             para_style_id = style_tag.attrib.get(f"{{{ns['w']}}}val")
#
#                         # 对齐方式
#                         jc = pPr.find("w:jc", ns)
#                         if jc is not None:
#                             alignment = jc.attrib.get(f"{{{ns['w']}}}val", "left")
#                         elif para_style_id and para_style_id in style_align_map:
#                             alignment = style_align_map[para_style_id]
#
#                         # 首行缩进
#                         ind = pPr.find("w:ind", ns)
#                         if ind is not None and f"{{{ns['w']}}}firstLine" in ind.attrib:
#                             first_line_indent = int(ind.attrib.get(f"{{{ns['w']}}}firstLine")) / 20  # 转 pt
#
#                     font_sizes = []
#                     for r in child.findall("w:r", ns):
#                         rPr = r.find("w:rPr", ns)
#                         if rPr is not None:
#                             # 字号
#                             sz = rPr.find("w:sz", ns)
#                             if sz is not None:
#                                 sz_val = int(sz.attrib.get(f"{{{ns['w']}}}val")) / 2
#                                 font_sizes.append(sz_val)
#
#                             # 字体
#                             fonts = rPr.find("w:rFonts", ns)
#                             if fonts is not None:
#                                 font_name = fonts.attrib.get(f"{{{ns['w']}}}ascii", font_name)
#                     # print(text.strip())
#                     # print(font_sizes)
#                     # 根据最多出现的字号选取
#                     if font_sizes:
#                         counter = Counter(font_sizes)
#                         most_common = counter.most_common()
#                         if most_common:  # 再次判断是否非空，确保稳健
#                             max_freq = most_common[0][1]
#                             # 取所有“频率等于最大频率”的字号
#                             most_frequent_sizes = [size for size, freq in most_common if freq == max_freq]
#                             # 从这些中选最大字号
#                             font_size = max(most_frequent_sizes)
#                         else:
#                             font_size = 12
#                     elif para_style_id and para_style_id in style_font_size_map:
#                         font_size = style_font_size_map[para_style_id]
#                     else:
#                         font_size = 12
#                     # print(font_size)
#
#                     # 拼接到结果中
#                     result.append({
#                         "type": "text",
#                         "content": text.strip(),
#                         "style": {
#                             "alignment": alignment,
#                             "first_line_indent_pt": first_line_indent,
#                             "font_name": font_name,
#                             "font_size_pt": font_size
#                         }
#                     })
#                     #print("*" * 30)
#
#                 # 查找 embedded 图片
#                 for blip in child.findall(".//a:blip", ns):
#                     r_embed = blip.attrib.get("{%s}embed" % ns["r"])
#                     if r_embed and r_embed in rels_map:
#                         image_path = "word/" + rels_map[r_embed]
#                         if image_path in docx_zip.namelist():
#                             image_data = docx_zip.read(image_path)
#                             b64 = base64.b64encode(image_data).decode()
#
#                             # 获取图片大小
#                             width_pt = height_pt = width_cm = height_cm = None
#                             for drawing in child.findall(".//w:drawing", ns):
#                                 extent = drawing.find(".//wp:extent", {
#                                     "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
#                                 })
#                                 if extent is not None:
#                                     cx = int(extent.attrib.get("cx", 0))
#                                     cy = int(extent.attrib.get("cy", 0))
#                                     width_pt = round(cx / 12700, 2)  # 转 pt
#                                     height_pt = round(cy / 12700, 2)
#                                     width_cm = round(cx / 360000, 2)
#                                     height_cm = round(cy / 360000, 2)
#                                     break  # 找到一个即可
#
#                             result.append({
#                                 "type": "image",
#                                 "filename": os.path.basename(image_path),
#                                 "content": b64,
#                                 "size": {
#                                     "width_pt": width_pt,
#                                     "height_pt": height_pt,
#                                     "width_cm": width_cm,
#                                     "height_cm": height_cm
#                                 }
#                             })
#
#             elif tag == "tbl":  # 表格
#                 rows = []
#                 for row in child.findall("w:tr", ns):
#                     cells = []
#                     for cell in row.findall("w:tc", ns):
#                         texts = [t.text for t in cell.findall(".//w:t", ns) if t.text]
#                         cells.append("".join(texts).strip())
#                     rows.append(cells)
#
#                 if table_markdown:
#                     try:
#                         rows = convert_table_to_markdown(rows)
#                     except Exception as e:
#                         pass
#                 result.append({"type": "table", "content": rows})
#
#     return result



def convert_table_to_markdown(table_data: list) -> str:
    """将二维列表转换为Markdown表格"""
    if not isinstance(table_data, list) or len(table_data) == 0:
        if isinstance(table_data, str):
            return table_data
        return ""

    markdown_lines = []
    header = table_data[0]
    if not isinstance(header, list):
        return ""

    markdown_lines.append("| " + " | ".join(str(cell) for cell in header) + " |")
    markdown_lines.append("| " + " | ".join("---" for _ in header) + " |")

    for row in table_data[1:]:
        if not isinstance(row, list):
            continue
        formatted_row = [str(row[i]) if i < len(row) else "" for i in range(len(header))]
        markdown_lines.append("| " + " | ".join(formatted_row) + " |")

    return "\n".join(markdown_lines)

class DocxGenerator:
    def __init__(self, template_path: str = None):
        #self.doc = Document()
        if template_path and os.path.exists(template_path):
            self.doc = Document(template_path)
            print(f"使用模板文件: {template_path}")
        else:
            self.doc = Document()
        self.default_font_name = "宋体"
        self.default_font_size = Pt(12)
        #self._set_default_font('宋体', 12)  # 默认字体和字号

    def _set_default_font(self, font_name: str, font_size: Pt):
        """设置文档默认字体"""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = font_name
        font.size = font_size

    def set_page_margins(self, left: Cm = 2, right: Cm = 2, top: Cm = 2, bottom: Cm = 2):
        """设置所有页面的边距（单位：厘米）"""
        for section in self.doc.sections:  # 遍历所有 sections
            section.left_margin = Cm(left)
            section.right_margin = Cm(right)
            section.top_margin = Cm(top)
            section.bottom_margin = Cm(bottom)

    ## 添加段落
    def add_paragraph(self, text: str = "", font_name: str = "宋体", font_size: Pt = 12, bold: bool = False,
                      color: str = "#000000", alignment: str = "left", line_spacing: float = 1.0,
                      space_after: float = 0, first_line_indent: float = 24,
                        bg_color: str = None):
        """添加段落并设置样式
        Args:
            first_line_indent: 首行缩进值（单位：字符）
            space_after: 段后间距（单位：磅）
        """
        para = self.doc.add_paragraph()

        # 设置首行缩进（1字符 ≈ 14.35磅）
        para.paragraph_format.first_line_indent = Pt(first_line_indent)

        # 文本对齐方式
        align_map = {
            'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
            'justify': WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        }
        para.alignment = align_map.get(alignment.lower(), WD_PARAGRAPH_ALIGNMENT.LEFT)

        # 段落格式
        para.paragraph_format.line_spacing = line_spacing
        para.paragraph_format.space_after = Pt(space_after)

        # 添加文本并设置样式
        run = para.add_run(text)

        # 👇 设置段落背景色（shading）
        if bg_color and bg_color.startswith("#"):
            hex_color = bg_color.lstrip("#")
            if len(hex_color) == 3:
                hex_color = "".join([c * 2 for c in hex_color])
            if len(hex_color) == 6:
                # 映射 HEX → Word 预设荧光笔颜色名
                color_map = {
                    "FFFF00": "yellow",  # 黄色
                    "00FFFF": "cyan",  # 青色
                    "FF0000": "red",  # 红色
                    "00FF00": "green",  # 绿色
                    "FF8000": "darkYellow",  # 橙色
                    "FF00FF": "pink",  # 粉红
                    "800080": "darkBlue",  # 紫色（近似）
                    "C0C0C0": "lightGray",  # 浅灰
                    "0000FF": "blue",  # 蓝色
                    "FFFFFF": "white",  # 清除高亮
                }
                word_color = color_map.get(hex_color.upper(), "yellow")  # 默认黄色
                highlight = OxmlElement('w:highlight')
                highlight.set(qn('w:val'), word_color)
                rPr = run._element.get_or_add_rPr()
                existing_highlight = rPr.find(qn('w:highlight'))
                if existing_highlight is not None:
                    rPr.remove(existing_highlight)
                rPr.append(highlight)

        # run.font.name = font_name
        # run.font.size = Pt(font_size)
        # run.font.bold = bold
        self._force_font(run, font_name, Pt(font_size), bold)

        # 处理十六进制颜色码
        if color.startswith("#"):
            hex_color = color.lstrip("#")
            if len(hex_color) == 3:  # 支持简写格式 #RGB → #RRGGBB
                hex_color = "".join([c * 2 for c in hex_color])
            rgb = tuple(int(hex_color[i:i + 2], 16) for i in (0, 2, 4))
            run.font.color.rgb = RGBColor(*rgb)
        else:
            run.font.color.rgb = RGBColor(0, 0, 0)  # 默认黑色

        return para

    # 添加表格
    def add_table_from_markdown(self, data, style: str = "Table Grid", alignment="center", font_name="宋体",font_size=12,cell_width: Optional[float] = None):
        if style in ["three_line","three line","Three Line Table"]:
            return self.add_three_line_table(data, font_name=font_name, font_size=font_size,
                                              alignment=alignment, cell_width=cell_width)
        try:
            # 支持 Markdown 格式
            if isinstance(data, str):
                # 拆行、去空白
                lines = [line.strip() for line in data.strip().split('\n') if line.strip()]
                if len(lines) < 2:
                    print("表格内容不足")
                    return

                # 第1行是表头，第2行是分隔线，之后是数据
                headers = [h.strip() for h in lines[0].strip('|').split('|')]
                data_lines = []
                for line in lines[1:]:
                    if re.match(r'^\s*\|?(?:\s*:?-+:?\s*\|)+\s*$', line):
                        continue  # 跳过分隔线
                    cells = [cell.strip() for cell in line.strip('|').split('|')]
                    data_lines.append(cells)

            # 支持 二维数组 格式
            elif isinstance(data, list) and all(isinstance(row, list) for row in data):
                if len(data) < 1:
                    print("空表格数据")
                    return
                headers = data[0]
                data_lines = data[1:]

            else:
                print("不支持的表格输入格式")
                return
            #print(data_lines)
            # 对齐方式映射
            align_map = {
                'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
                'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
                'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
                'justify': WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            }

            # 创建表格
            table = self.doc.add_table(rows=len(data_lines) + 1, cols=len(headers))
            if style and style != "":
                table.style = style
            table.alignment = align_map.get(alignment.lower(), WD_PARAGRAPH_ALIGNMENT.LEFT)

            # 写表头
            for col_idx, header in enumerate(headers):
                cell = table.cell(0, col_idx)
                cell.text = ''
                p = cell.paragraphs[0]
                p.clear()
                run = p.add_run(header)
                run.font.bold = True
                run.font.name = font_name
                run.font.size = Pt(font_size)
                run.font.color.rgb = RGBColor(0, 0, 0)
                p.alignment = align_map.get(alignment.lower(), WD_PARAGRAPH_ALIGNMENT.LEFT)

            # 表格内容
            for row_idx, row_cells in enumerate(data_lines, start=1):
                for col_idx in range(min(len(row_cells), len(headers))):  # 避免越界
                    cell = table.cell(row_idx, col_idx)
                    cell.text = ''
                    p = cell.paragraphs[0]
                    p.clear()
                    run = p.add_run(str(row_cells[col_idx]))
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    p.alignment = align_map.get(alignment.lower(), WD_PARAGRAPH_ALIGNMENT.LEFT)

            # 自动设置列宽
            # for col in table.columns:
            #     for cell in col.cells:
            #         cell.width = Inches(1.5)
            # ✅ 设置列宽（如果用户指定了宽度）
            if cell_width and cell_width > 0:
                table.autofit = False
                for column in table.columns:
                    for cell in column.cells:
                            cell.width = Inches(cell_width)
        except Exception as e:
            print(f"表格写入失败: {str(e)}")

    def _set_cell_border(self, cell, **kwargs):
        """
        内部辅助函数：设置单元格边框
        Usage: self._set_cell_border(cell, top={"sz": 12, "val": "single", "color": "#000000"})
        sz 单位为 1/8 pt, 12 表示 1.5pt
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        for edge, params in kwargs.items():
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, val in params.items():
                element.set(qn('w:{}'.format(key)), str(val))

    def add_three_line_table(self, data, alignment="center", font_name="宋体", font_size=12,
                             cell_width: Optional[float] = None):
        """新增功能：添加学术论文标准三线表"""
        #print("三线表")
        try:
            if isinstance(data, str):
                lines = [line.strip() for line in data.strip().split('\n') if line.strip()]
                if len(lines) < 2: return
                headers = [h.strip() for h in lines[0].strip('|').split('|')]
                data_lines = []
                for line in lines[1:]:
                    if re.match(r'^\s*\|?(?:\s*:?-+:?\s*\|)+\s*$', line): continue
                    data_lines.append([cell.strip() for cell in line.strip('|').split('|')])
            elif isinstance(data, list) and all(isinstance(row, list) for row in data):
                headers = data[0]
                data_lines = data[1:]
            else:
                return

            align_map = {'left': WD_PARAGRAPH_ALIGNMENT.LEFT, 'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
                         'right': WD_PARAGRAPH_ALIGNMENT.RIGHT}
            table = self.doc.add_table(rows=len(data_lines) + 1, cols=len(headers))
            table.style = None  # 三线表不使用预设样式
            table.alignment = align_map.get(alignment.lower(), WD_PARAGRAPH_ALIGNMENT.LEFT)

            # 边框规格定义
            border_thick = {"sz": 12, "val": "single", "color": "#000000"}  # 1.5pt
            border_thin = {"sz": 6, "val": "single", "color": "#000000"}  # 0.75pt
            border_none = {"val": "nil"}

            for r_idx in range(len(data_lines) + 1):
                row_data = headers if r_idx == 0 else data_lines[r_idx - 1]
                for c_idx in range(len(headers)):
                    cell = table.cell(r_idx, c_idx)
                    # 写入内容
                    p = cell.paragraphs[0]
                    p.clear()
                    txt = str(row_data[c_idx]) if c_idx < len(row_data) else ""
                    run = p.add_run(txt)
                    run.font.name, run.font.size = font_name, Pt(font_size)
                    if r_idx == 0: run.font.bold = True
                    p.alignment = align_map.get(alignment.lower(), WD_PARAGRAPH_ALIGNMENT.LEFT)

                    # 设置边框逻辑
                    borders = {"left": border_none, "right": border_none, "insideV": border_none,
                               "insideH": border_none}
                    if r_idx == 0:  # 顶线与表头底线
                        borders["top"] = border_thick
                        borders["bottom"] = border_thin
                    elif r_idx == len(data_lines):  # 底部封底线
                        borders["bottom"] = border_thick
                        borders["top"] = border_none
                    else:  # 中间行无横线
                        borders["top"] = border_none
                        borders["bottom"] = border_none

                    self._set_cell_border(cell, **borders)

            if cell_width and cell_width > 0:
                table.autofit = False
                for col in table.columns:
                    for cell in col.cells: cell.width = Inches(cell_width)
        except Exception as e:
            print(f"三线表写入失败: {str(e)}")

    def add_table_from_markdown2(self, data, style: str = "Table Grid", alignment="center", font_name="宋体",
                                font_size=12,
                                cell_width: Optional[float] = None, column_alignments: Optional[list] = None,
                                bold_content: bool = False, text_color: str = "#000000", vertical_align: str = "center"):
        try:
            # 支持 Markdown 格式
            if isinstance(data, str):
                # 拆行、去空白
                lines = [line.strip() for line in data.strip().split('\n') if line.strip()]
                if len(lines) < 2:
                    print("表格内容不足")
                    return

                # 第1行是表头，第2行是分隔线，之后是数据
                headers = [h.strip() for h in lines[0].strip('|').split('|')]
                data_lines = []
                for line in lines[1:]:
                    if re.match(r'^\s*\|?(?:\s*:?-+:?\s*\|)+\s*$', line):
                        continue  # 跳过分隔线
                    cells = [cell.strip() for cell in line.strip('|').split('|')]
                    data_lines.append(cells)

            # 支持 二维数组 格式
            elif isinstance(data, list) and all(isinstance(row, list) for row in data):
                if len(data) < 1:
                    print("空表格数据")
                    return
                headers = data[0]
                data_lines = data[1:]

            else:
                print("不支持的表格输入格式")
                return

            # 对齐方式映射
            align_map = {
                'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
                'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
                'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
                'justify': WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            }

            # 创建表格
            table = self.doc.add_table(rows=len(data_lines) + 1, cols=len(headers))
            if style and style != "":
                table.style = style
            table.alignment = align_map.get(alignment.lower(), WD_PARAGRAPH_ALIGNMENT.LEFT)

            # 处理颜色
            if text_color.startswith("#"):
                hex_color = text_color.lstrip("#")
                if len(hex_color) == 3:
                    hex_color = "".join([c * 2 for c in hex_color])
                rgb_color = tuple(int(hex_color[i:i + 2], 16) for i in (0, 2, 4))
                font_color = RGBColor(*rgb_color)
            else:
                font_color = RGBColor(0, 0, 0)  # 默认黑色

            # 写表头
            for col_idx, header in enumerate(headers):
                cell = table.cell(0, col_idx)
                cell.text = ''
                p = cell.paragraphs[0]
                p.clear()

                # 设置垂直对齐 - 使用参数值
                tcPr = cell._tc.get_or_add_tcPr()
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), vertical_align)
                tcPr.append(vAlign)

                run = p.add_run(header)
                run.font.bold = True  # 表头默认加粗
                run.font.name = font_name
                run.font.size = Pt(font_size)
                run.font.color.rgb = font_color

                # 设置列对齐方式
                if column_alignments and col_idx < len(column_alignments):
                    col_alignment = column_alignments[col_idx]
                    p.alignment = align_map.get(col_alignment.lower(), WD_PARAGRAPH_ALIGNMENT.LEFT)
                else:
                    p.alignment = align_map.get(alignment.lower(), WD_PARAGRAPH_ALIGNMENT.LEFT)

            # 表格内容
            for row_idx, row_cells in enumerate(data_lines, start=1):
                for col_idx in range(min(len(row_cells), len(headers))):  # 避免越界
                    cell = table.cell(row_idx, col_idx)
                    cell.text = ''
                    p = cell.paragraphs[0]
                    p.clear()

                    # 设置垂直对齐 - 使用参数值
                    tcPr = cell._tc.get_or_add_tcPr()
                    vAlign = OxmlElement('w:vAlign')
                    vAlign.set(qn('w:val'), vertical_align)
                    tcPr.append(vAlign)

                    run = p.add_run(str(row_cells[col_idx]))
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
                    run.font.color.rgb = font_color
                    run.font.bold = bold_content  # 内容是否加粗

                    # 设置列对齐方式
                    if column_alignments and col_idx < len(column_alignments):
                        col_alignment = column_alignments[col_idx]
                        p.alignment = align_map.get(col_alignment.lower(), WD_PARAGRAPH_ALIGNMENT.LEFT)
                    else:
                        p.alignment = align_map.get(alignment.lower(), WD_PARAGRAPH_ALIGNMENT.LEFT)

            # 设置列宽
            if cell_width and cell_width > 0:
                table.autofit = False
                for column in table.columns:
                    for cell in column.cells:
                        cell.width = Inches(cell_width)

        except Exception as e:
            print(f"表格写入失败: {str(e)}")

    def add_image(self, image_data: str, width: float = None,  height: float = None,  alignment: str = "center"):
        try:
            # 判断图片类型
            if image_data.strip().startswith(("http://", "https://")):
                # 网络图片，使用 httpx 获取
                with httpx.Client(timeout=30.0) as client:
                    # 设置更细致的超时：连接5秒，读写60秒
                    timeout_config = httpx.Timeout(60.0, connect=10.0)

                    # 增加简单的重试逻辑 (最多尝试3次)
                    for attempt in range(3):
                        try:
                            with httpx.Client(timeout=timeout_config, follow_redirects=True) as client:
                                response = client.get(image_data)
                                response.raise_for_status()
                                image_bytes = response.content

                                #print(f"{attempt}次成功")
                                break  # 下载成功，跳出循环
                        except (httpx.TimeoutException, httpx.NetworkError) as net_err:
                            if attempt == 2:  # 最后一次尝试仍失败
                                raise net_err
                            time.sleep(1)  # 等待1秒后重试

            elif "base64," in image_data:
                # base64 图片（可能是 data:image/png;base64,...）
                base64_str = image_data.split("base64,", 1)[-1]
                image_bytes = base64.b64decode(base64_str)
            else:
                raise ValueError("不支持的图片格式")

            # 1. 下载完成后，先用 Pillow 处理图片
            temp_stream = BytesIO(image_bytes)

            # 使用 Pillow 打开，处理 Adobe/CMYK 兼容性
            with Image.open(temp_stream) as img:
                # 判断：如果已经是 PNG，且模式不是 CMYK，则不处理直接使用
                if img.format == 'PNG' and img.mode != 'CMYK':
                    # 直接回滚指针，复用原始流（保留透明度）
                    temp_stream.seek(0)
                    output_stream = temp_stream
                else:
                    # 如果是 JPG (特别是 Adobe CMYK) 或其他格式，执行转换
                    # 转换逻辑：非 RGB 模式统一转 RGB
                    if img.mode != 'RGB':
                        # 注意：如果是从透明格式转 RGB，Pillow 默认会变黑底
                        # 但因为上面判断了 PNG 不走这里，所以这里的 img 主要是针对 JPG/CMYK
                        img = img.convert('RGB')

                    output_stream = BytesIO()
                    img.save(output_stream, format='PNG')  # 统一转为 Word 兼容性最好的 PNG 封装
                    output_stream.seek(0)

            # 添加段落
            paragraph = self.doc.add_paragraph()
            align_map = {
                'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
                'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
                'right': WD_PARAGRAPH_ALIGNMENT.RIGHT
            }
            paragraph.alignment = align_map.get(alignment.lower(), WD_PARAGRAPH_ALIGNMENT.CENTER)

            # 插入图片
            run = paragraph.add_run()

            if width and height:
                run.add_picture(output_stream, width=Cm(width), height=Cm(height))
            elif width:
                run.add_picture(output_stream, width=Cm(width))
            elif height:
                run.add_picture(output_stream, height=Cm(height))
            else:
                run.add_picture(output_stream)

            output_stream.close()
            temp_stream.close()

            return paragraph
        except Exception as e:
            #error_detail = traceback.format_exc()
            paragraph = self.add_paragraph(f"添加图片失败: {e}",color="#FF0000")
            #print(f"完整错误详情:\n{error_detail}")
            return paragraph

    def add_float_image_absolute(self, image_data: str, width: float = None, height: float = None,
                                 pos_x: float = 0, pos_y: float = 0, wrap_type: str = "topBottom"):
        """
        添加“上下型环绕”图片
        dist_t, dist_b: 图片距离上方和下方文字的间距（单位：磅）
        """
        try:
            # 判断图片类型
            if image_data.strip().startswith(("http://", "https://")):
                # 网络图片，使用 httpx 获取
                with httpx.Client(timeout=30.0) as client:
                    # 设置更细致的超时：连接5秒，读写60秒
                    timeout_config = httpx.Timeout(60.0, connect=10.0)

                    # 增加简单的重试逻辑 (最多尝试3次)
                    for attempt in range(3):
                        try:
                            with httpx.Client(timeout=timeout_config, follow_redirects=True) as client:
                                response = client.get(image_data)
                                response.raise_for_status()
                                image_bytes = response.content

                                # print(f"{attempt}次成功")
                                break  # 下载成功，跳出循环
                        except (httpx.TimeoutException, httpx.NetworkError) as net_err:
                            if attempt == 2:  # 最后一次尝试仍失败
                                raise net_err
                            time.sleep(1)  # 等待1秒后重试

            elif "base64," in image_data:
                # base64 图片（可能是 data:image/png;base64,...）
                base64_str = image_data.split("base64,", 1)[-1]
                image_bytes = base64.b64decode(base64_str)
            else:
                raise ValueError("不支持的图片格式")

            # 1. 下载完成后，先用 Pillow 处理图片
            temp_stream = BytesIO(image_bytes)

            # 使用 Pillow 打开，处理 Adobe/CMYK 兼容性
            with Image.open(temp_stream) as img:
                # 判断：如果已经是 PNG，且模式不是 CMYK，则不处理直接使用
                if img.format == 'PNG' and img.mode != 'CMYK':
                    # 直接回滚指针，复用原始流（保留透明度）
                    temp_stream.seek(0)
                    output_stream = temp_stream
                else:
                    # 如果是 JPG (特别是 Adobe CMYK) 或其他格式，执行转换
                    # 转换逻辑：非 RGB 模式统一转 RGB
                    if img.mode != 'RGB':
                        # 注意：如果是从透明格式转 RGB，Pillow 默认会变黑底
                        # 但因为上面判断了 PNG 不走这里，所以这里的 img 主要是针对 JPG/CMYK
                        img = img.convert('RGB')

                    output_stream = BytesIO()
                    img.save(output_stream, format='PNG')  # 统一转为 Word 兼容性最好的 PNG 封装
                    output_stream.seek(0)

            # 2. 插入图片
            paragraph = self.doc.add_paragraph()
            run = paragraph.add_run()
            picture = run.add_picture(output_stream, width=Cm(width) if width else None,
                                      height=Cm(height) if height else None)

            inline = picture._inline
            cx = inline.extent.cx
            cy = inline.extent.cy

            # 3. 构造绝对定位 XML
            # relativeFrom="page" 是实现 x=0, y=0 的关键
            # wp:wrapTopAndBottom (上下型) 或 wp:wrapNone (浮于文字上方)
            wrap_xml = '<wp:wrapTopAndBottom/>' if wrap_type == "topBottom" else '<wp:wrapNone/>'

            x_emu = int(pos_x * 360000)
            y_emu = int(pos_y * 360000)

            anchor_xml = f"""
                <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" \
                relativeHeight="251658240" behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1" {nsdecls('wp', 'r', 'a', 'pic')}>
                    <wp:simplePos x="0" y="0"/>
                    <wp:positionH relativeFrom="page">
                        <wp:posOffset>{x_emu}</wp:posOffset>
                    </wp:positionH>
                    <wp:positionV relativeFrom="page">
                        <wp:posOffset>{y_emu}</wp:posOffset>
                    </wp:positionV>
                    <wp:extent cx="{cx}" cy="{cy}"/>
                    {wrap_xml}
                    <wp:docPr id="1" name="Picture 1"/>
                    <wp:cNvGraphicFramePr>
                        <a:graphicFrameLocks xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" noChangeAspect="1"/>
                    </wp:cNvGraphicFramePr>
                    <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">
                        <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
                            {inline.graphic.graphicData.xml}
                        </a:graphicData>
                    </a:graphic>
                </wp:anchor>
                """
            anchor_element = parse_xml(anchor_xml)
            inline.getparent().replace(inline, anchor_element)

        except Exception as e:
            self.add_paragraph(f"设置上下型环绕失败: {e}", color="#FF0000")

    def add_paragraph_comment(self,para,comment_text = "批注",author="ay-office-word",initials=""):
        # 添加基础批注
        comment = self.doc.add_comment(
            runs=para.runs[0],  # 绑定到段落的第一个Run
            text=comment_text,
            author=author,
            initials=initials
        )
        return comment

    def add_line_break(self,para = None):
        """添加换行"""
        if para is not None:
            try:
                para.add_run().add_break()
            except Exception as e:
                self.doc.add_paragraph()
        else:
            self.doc.add_paragraph()

    def add_page_break(self):
        """添加分页符"""
        self.doc.add_page_break()

    def add_heading(self, text: str, level: int = 1, style: str = None):
        """添加标题（支持自定义样式）"""
        heading = self.doc.add_heading(level=level)
        run = heading.add_run(text)
        if style == 'modern':
            run.font.name = 'Calibri Light'
            run.font.size = Pt(18 if level == 1 else 16 if level == 2 else 14)
            run.font.color.rgb = RGBColor(44, 62, 80)

    ''' 强制设置字体 '''
    def _force_font(self, run, font_name=None, font_size=None, bold=None, color_rgb=None):
        """强制给 run 设置中英文字体"""
        font_name = font_name or self.default_font_name
        font_size = font_size or self.default_font_size

        run.font.name = font_name  # 英文字体
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)  # 中文字体
        run.font.size = font_size
        if bold is not None:
            run.font.bold = bold
        if color_rgb:
            run.font.color.rgb = RGBColor(*color_rgb)

    def enable_line_numbering(self, start=1, count_by=1, restart='continuous'):
        """
        启用行号
        Args:
            start: 行号起始值
            count_by: 行号间隔（每多少行编号一次）
            restart: 'continuous' 连续编号, 'newPage' 每页重新开始, 'newSection' 每节重新开始
        """
        section = self.doc.sections[0]
        sectPr = section._sectPr
        lnNumType_xml = (
            f'<w:lnNumType xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            f' w:countBy="{count_by}" w:restart="{restart}"/>'
        )
        sectPr.append(parse_xml(lnNumType_xml))


    def add_page_number(self):
        section = self.doc.sections[0]
        footer = section.footer

        for element in footer._element.xpath(".//w:p"):
            footer._element.remove(element)

        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        run._r.append(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)


    def save(self, filename: str):
        """保存文档"""
        self.doc.save(filename)

    def buildBytes(self):
        # 保存到字节流
        """将文档转换为字节流（自动处理资源释放）"""
        byte_io = BytesIO()
        self.doc.save(byte_io)
        byte_io.seek(0)
        docx_bytes = byte_io.getvalue()  # 使用 getvalue() 替代 read()
        byte_io.close()  # 显式关闭
        return docx_bytes

@contextmanager
def no_proxy():
    original_http = os.environ.get('HTTP_PROXY')
    original_https = os.environ.get('HTTPS_PROXY')
    original_no = os.environ.get('NO_PROXY')

    try:
        #print("临时禁用所有代理")
        os.environ.update({
            'HTTP_PROXY': '',
            'HTTPS_PROXY': '',
            'NO_PROXY': '*',  # 禁用所有代理
            'http_proxy': '',  # 小写变量也清空
            'https_proxy': '',
            'no_proxy': '*'
        })
        yield
    finally:
        # 恢复原始环境变量
        if original_http is not None:
            os.environ['HTTP_PROXY'] = original_http
        if original_https is not None:
            os.environ['HTTPS_PROXY'] = original_https
        if original_no is not None:
            os.environ['NO_PROXY'] = original_no
